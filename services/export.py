from __future__ import annotations

from datetime import datetime
from io import BytesIO

from docx import Document

from core.models import TaskRequest


def build_markdown(request: TaskRequest, answer: str, sources: str = "") -> str:
    parts = [
        "# 粤见非遗生成结果",
        "",
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "## 原始需求",
        request.raw_request.strip(),
        "",
        answer.strip(),
    ]
    if sources.strip():
        parts.extend(["", sources.strip()])
    return "\n".join(parts).strip() + "\n"


def build_plain_text(request: TaskRequest, answer: str, sources: str = "") -> str:
    markdown = build_markdown(request, answer, sources)
    return markdown.replace("# ", "").replace("## ", "").replace("### ", "")


def build_docx(request: TaskRequest, answer: str, sources: str = "") -> bytes:
    document = Document()
    document.add_heading("粤见非遗生成结果", level=0)
    document.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_heading("原始需求", level=1)
    document.add_paragraph(request.raw_request.strip())

    for line in answer.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped[4:], level=2)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=1)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:], style="List Bullet")
        else:
            document.add_paragraph(stripped)

    if sources.strip():
        document.add_heading("本次检索来源", level=1)
        for line in sources.splitlines():
            stripped = line.strip()
            if stripped.startswith("- "):
                document.add_paragraph(stripped[2:], style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
